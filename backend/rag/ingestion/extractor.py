"""Extract text from 3GPP spec ZIP files containing DOCX documents."""

import zipfile
from pathlib import Path
from typing import Generator
from docx import Document
from loguru import logger


def iter_spec_zips(specs_dir: str, releases: list[str], series: list[str]) -> Generator[Path, None, None]:
    """Yield ZIP file paths matching the given releases and series."""
    base = Path(specs_dir)
    for release in releases:
        for serie in series:
            folder = base / release / serie
            if not folder.exists():
                logger.warning(f"Folder not found: {folder}")
                continue
            for zip_path in sorted(folder.glob("*.zip")):
                yield zip_path


def extract_text_from_docx(docx_path_or_bytes) -> str:
    """Extract plain text from a DOCX file or bytes object."""
    try:
        doc = Document(docx_path_or_bytes)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"Failed to parse DOCX: {e}")
        return ""


def extract_zip(zip_path: Path) -> Generator[tuple[str, str], None, None]:
    """
    Yield (filename, text) tuples from all DOCX files inside a ZIP.
    Skips cover docs (usually very short).
    """
    try:
        with zipfile.ZipFile(zip_path, "r") as zf:
            docx_files = [f for f in zf.namelist() if f.endswith(".docx")]
            for docx_name in docx_files:
                try:
                    with zf.open(docx_name) as f:
                        import io
                        text = extract_text_from_docx(io.BytesIO(f.read()))
                        if len(text) > 200:  # skip near-empty docs
                            yield docx_name, text
                except Exception as e:
                    logger.warning(f"Skipping {docx_name} in {zip_path.name}: {e}")
    except Exception as e:
        logger.error(f"Failed to open ZIP {zip_path}: {e}")


def parse_spec_metadata(zip_path: Path) -> dict:
    """Extract release, series, spec number from file path."""
    parts = zip_path.parts
    release = next((p for p in parts if p.startswith("Rel-")), "unknown")
    series = next((p for p in parts if p.endswith("_series")), "unknown")
    spec_name = zip_path.stem  # e.g. "38101-1-i70"
    spec_number = spec_name.split("-")[0] if "-" in spec_name else spec_name
    return {
        "release": release,
        "series": series,
        "spec_number": spec_number,
        "spec_name": spec_name,
        "zip_file": zip_path.name,
    }
